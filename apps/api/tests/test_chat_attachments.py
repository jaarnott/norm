"""Uploading documents to the chat and sending them to the model.

Covers turning an uploaded file into the right Anthropic content block (native
for PDF/image, extracted text for Office/text), linking an upload to its thread,
injecting the attaching turn as a block list, and the get_attachment tool the
model uses to re-open a file on a later turn.
"""

import io
import uuid

import pytest

from app.agents import context_builder as CB
from app.agents import internal_tools as IT
from app.db.models import Message, UploadedDocument
from app.services import attachments as A
from tests.conftest import _make_thread


def _docx_bytes(text: str) -> bytes:
    import docx

    d = docx.Document()
    d.add_paragraph(text)
    b = io.BytesIO()
    d.save(b)
    return b.getvalue()


def _xlsx_bytes(rows) -> bytes:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    for r in rows:
        ws.append(r)
    b = io.BytesIO()
    wb.save(b)
    return b.getvalue()


def _upload(db, user, *, content_type, filename, data=b"bytes", thread_id=None):
    doc = UploadedDocument(
        id=str(uuid.uuid4()),
        user_id=user.id,
        filename=filename,
        content_type=content_type,
        size=len(data),
        data=data,
        thread_id=thread_id,
    )
    db.add(doc)
    db.flush()
    return doc


class TestBuildContentBlock:
    def test_pdf_is_a_native_document_block(self):
        blk = A.build_content_block(b"%PDF-1.4 x", "application/pdf", "a.pdf")
        assert blk["type"] == "document"
        assert blk["source"]["media_type"] == "application/pdf"
        assert blk["source"]["type"] == "base64" and blk["source"]["data"]

    def test_png_is_a_native_image_block(self):
        blk = A.build_content_block(b"\x89PNG\r\n", "image/png", "a.png")
        assert blk["type"] == "image"
        assert blk["source"]["media_type"] == "image/png"

    def test_plain_text_is_extracted(self):
        blk = A.build_content_block(b"hello,world\n1,2", "text/csv", "d.csv")
        assert blk["type"] == "text"
        assert "hello,world" in blk["text"]
        assert "[Attachment: d.csv]" in blk["text"]

    def test_docx_is_extracted_to_text(self):
        blk = A.build_content_block(
            _docx_bytes("Quarterly plan for the bar team"),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "plan.docx",
        )
        assert blk["type"] == "text"
        assert "Quarterly plan for the bar team" in blk["text"]

    def test_xlsx_is_extracted_to_text(self):
        blk = A.build_content_block(
            _xlsx_bytes([["Item", "Qty"], ["Gin", 5]]),
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "stock.xlsx",
        )
        assert blk["type"] == "text"
        assert "Item" in blk["text"] and "Gin" in blk["text"]

    def test_extension_fallback_when_content_type_is_generic(self):
        # Browsers sometimes send application/octet-stream — fall back to the name.
        blk = A.build_content_block(
            _docx_bytes("from the extension"),
            "application/octet-stream",
            "notes.docx",
        )
        assert blk["type"] == "text" and "from the extension" in blk["text"]

    def test_unsupported_type_raises(self):
        with pytest.raises(ValueError):
            A.build_content_block(b"\x00\x01", "application/x-msdownload", "a.exe")


def _jpeg_bytes(w, h, *, quality=95, exif_orientation=None):
    import os

    from PIL import Image

    img = Image.frombytes("RGB", (w, h), os.urandom(w * h * 3))
    buf = io.BytesIO()
    kw = {"format": "JPEG", "quality": quality}
    if exif_orientation is not None:
        exif = img.getexif()
        exif[0x0112] = exif_orientation
        kw["exif"] = exif
    img.save(buf, **kw)
    return buf.getvalue()


class TestImageNormalization:
    """A large phone photo must be shrunk to fit Anthropic's per-image limits
    before it becomes a block, or the whole Messages request fails (prod thread
    eb83d8f0: a 7.78MB JPEG killed the turn before any model call)."""

    def test_large_image_is_downscaled_under_limits(self):
        from PIL import Image

        big = _jpeg_bytes(3000, 2000)
        assert len(big) > 5_000_000  # the failing shape
        out, media = A.normalize_image_for_anthropic(big, "image/jpeg")
        assert len(out) < 5_000_000
        assert max(Image.open(io.BytesIO(out)).size) <= 1568
        assert media == "image/jpeg"

    def test_small_in_bounds_image_is_untouched(self):
        small = _jpeg_bytes(800, 600)
        out, media = A.normalize_image_for_anthropic(small, "image/jpeg")
        assert out == small  # no needless recompression
        assert media == "image/jpeg"

    def test_exif_orientation_is_baked_in(self):
        from PIL import Image

        rotated = _jpeg_bytes(400, 200, exif_orientation=6)  # 90° CW
        out, _ = A.normalize_image_for_anthropic(rotated, "image/jpeg")
        assert out != rotated
        im = Image.open(io.BytesIO(out))
        assert im.size == (200, 400)  # dimensions swapped → upright
        assert im.getexif().get(0x0112, 1) == 1  # orientation cleared

    def test_corrupt_image_returns_original_without_raising(self):
        junk = b"\x00\x01 not an image"
        out, media = A.normalize_image_for_anthropic(junk, "image/jpeg")
        assert out == junk  # degrades gracefully, never crashes the turn

    def test_build_content_block_shrinks_a_large_image(self):
        import base64

        blk = A.build_content_block(_jpeg_bytes(3000, 2000), "image/jpeg", "p.jpg")
        assert blk["type"] == "image"
        assert len(base64.b64decode(blk["source"]["data"])) < 5_000_000


class TestLinkAndManifest:
    def test_link_stamps_thread_and_returns_refs(self, db_session, admin_user):
        thread = _make_thread(db_session, admin_user)
        d1 = _upload(
            db_session, admin_user, content_type="application/pdf", filename="a.pdf"
        )
        d2 = _upload(db_session, admin_user, content_type="image/png", filename="b.png")
        refs = A.link_chat_attachments(
            [d1.id, d2.id], thread.id, admin_user.id, db_session
        )
        assert [r["upload_id"] for r in refs] == [d1.id, d2.id]
        assert [r["filename"] for r in refs] == ["a.pdf", "b.png"]
        # link_chat_attachments stamps the thread on the (identity-mapped) row;
        # the caller (base.py) commits it.
        assert d1.thread_id == thread.id

    def test_link_drops_another_users_upload(
        self, db_session, admin_user, manager_user
    ):
        thread = _make_thread(db_session, admin_user)
        mine = _upload(
            db_session, admin_user, content_type="application/pdf", filename="mine.pdf"
        )
        theirs = _upload(
            db_session,
            manager_user,
            content_type="application/pdf",
            filename="theirs.pdf",
        )
        refs = A.link_chat_attachments(
            [mine.id, theirs.id], thread.id, admin_user.id, db_session
        )
        assert [r["upload_id"] for r in refs] == [mine.id]

    def test_manifest_lists_thread_attachments(self, db_session, admin_user):
        thread = _make_thread(db_session, admin_user)
        d = _upload(
            db_session,
            admin_user,
            content_type="application/pdf",
            filename="invoice.pdf",
            thread_id=thread.id,
        )
        db_session.add(
            Message(
                thread_id=thread.id,
                role="user",
                content="here",
                attachments=[
                    {
                        "upload_id": d.id,
                        "filename": "invoice.pdf",
                        "content_type": "application/pdf",
                    }
                ],
            )
        )
        db_session.flush()
        man = A.attachment_manifest(thread, db_session)
        assert man and "invoice.pdf" in man and d.id in man

    def test_manifest_none_when_no_attachments(self, db_session, admin_user):
        thread = _make_thread(db_session, admin_user)
        assert A.attachment_manifest(thread, db_session) is None


class TestContextInjection:
    def test_attaching_turn_becomes_a_block_list(self, db_session, admin_user):
        thread = _make_thread(db_session, admin_user)
        doc = _upload(
            db_session,
            admin_user,
            content_type="application/pdf",
            filename="x.pdf",
            thread_id=thread.id,
        )
        msg = Message(
            thread_id=thread.id,
            role="user",
            content="what does this say?",
            attachments=[
                {
                    "upload_id": doc.id,
                    "filename": "x.pdf",
                    "content_type": "application/pdf",
                }
            ],
        )
        db_session.add(msg)
        db_session.flush()

        built = CB.build_conversation_messages(
            [msg], "what does this say?", thread=thread, db=db_session
        )
        last = built[-1]
        assert last["role"] == "user"
        assert isinstance(last["content"], list)
        kinds = [b["type"] for b in last["content"]]
        assert "document" in kinds and kinds[-1] == "text"  # text block goes last
        # The manifest is inside the trailing text block.
        assert "get_attachment" in last["content"][-1]["text"]

    def test_no_attachments_stays_a_plain_string(self, db_session, admin_user):
        thread = _make_thread(db_session, admin_user)
        msg = Message(thread_id=thread.id, role="user", content="hi")
        db_session.add(msg)
        db_session.flush()
        built = CB.build_conversation_messages(
            [msg], "hi", thread=thread, db=db_session
        )
        assert isinstance(built[-1]["content"], str)


class TestEnsureAlternation:
    def test_list_content_merges_without_error(self):
        merged = CB._ensure_alternation(
            [
                {"role": "user", "content": [{"type": "text", "text": "a"}]},
                {"role": "user", "content": "b"},
            ]
        )
        assert len(merged) == 1
        assert isinstance(merged[0]["content"], list)
        assert {b["text"] for b in merged[0]["content"]} == {"a", "b"}


class TestAttachmentAwareRouting:
    """A file only reaches the model on the agent tool-loop path, so routing to
    the no-tool-loop "meta"/"unknown" help reply silently drops it. When the
    turn carries an attachment the router is told to pick a specialist instead.
    """

    def _router_system(self, db_session, monkeypatch, *, has_attachments):
        import anthropic

        import app.services.agent_config_service as acs
        from app.agents import router

        monkeypatch.setattr(
            acs, "get_system_prompt", lambda name, cdb: "ROUTER {domains}\nReturn JSON."
        )

        captured = {}

        class _Usage:
            input_tokens = 1
            output_tokens = 1

        class _Block:
            text = '{"domain": "executive_chef"}'

        class _Resp:
            content = [_Block()]
            usage = _Usage()

        class _Msgs:
            def create(self, **kw):
                captured["system"] = kw["system"]
                return _Resp()

        class _Client:
            def __init__(self, *a, **k):
                self.messages = _Msgs()

        monkeypatch.setattr(anthropic, "Anthropic", _Client)

        router.classify(
            "here is a document",
            ["procurement", "executive_chef", "reports"],
            db=db_session,
            config_db=db_session,
            has_attachments=has_attachments,
        )
        return captured.get("system", "")

    def test_guidance_is_injected_when_a_file_is_attached(
        self, db_session, monkeypatch
    ):
        s = self._router_system(db_session, monkeypatch, has_attachments=True)
        assert "ATTACHED" in s
        assert "meta" in s and "unknown" in s

    def test_no_guidance_without_an_attachment(self, db_session, monkeypatch):
        s = self._router_system(db_session, monkeypatch, has_attachments=False)
        assert "ATTACHED" not in s


class TestGetAttachmentTool:
    def test_returns_document_block_for_pdf(self, db_session, admin_user):
        thread = _make_thread(db_session, admin_user)
        doc = _upload(
            db_session,
            admin_user,
            content_type="application/pdf",
            filename="a.pdf",
            thread_id=thread.id,
        )
        out = IT._get_attachment({"attachment_id": doc.id}, db_session, thread.id)
        assert out["success"] is True
        assert out["_document"]["type"] == "document"

    def test_returns_text_for_office(self, db_session, admin_user):
        thread = _make_thread(db_session, admin_user)
        doc = _upload(
            db_session,
            admin_user,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename="p.docx",
            data=_docx_bytes("the extracted body"),
            thread_id=thread.id,
        )
        out = IT._get_attachment({"attachment_id": doc.id}, db_session, thread.id)
        assert out["success"] is True
        assert "_document" not in out
        assert "the extracted body" in out["data"]["text"]

    def test_refuses_another_threads_attachment(self, db_session, admin_user):
        t1 = _make_thread(db_session, admin_user)
        t2 = _make_thread(db_session, admin_user)
        doc = _upload(
            db_session,
            admin_user,
            content_type="application/pdf",
            filename="a.pdf",
            thread_id=t1.id,
        )
        out = IT._get_attachment({"attachment_id": doc.id}, db_session, t2.id)
        assert out["success"] is False

    def test_missing_id(self, db_session, admin_user):
        thread = _make_thread(db_session, admin_user)
        out = IT._get_attachment({}, db_session, thread.id)
        assert out["success"] is False
